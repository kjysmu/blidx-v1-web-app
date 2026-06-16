from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/kjysmu/Documents/Blidx/Blidx_V1_Web_App_Implementation_Report.docx")
PURPLE = "5548B8"
PURPLE_LIGHT = "EEEAFE"
INK = "202025"
MUTED = "66636D"
LINE = "DADCE0"
GREEN = "237A57"
GREEN_LIGHT = "E7F5ED"
BLUE_LIGHT = "EAF1FF"
RED_LIGHT = "FFF1F1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, label, text, fill=PURPLE_LIGHT, color=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_table(doc):
    rows = [
        ("Web application shell", "Working", "Responsive desktop/mobile navigation and live screens"),
        ("Content Bank", "Working", "Entries persist locally and are reused in draft generation"),
        ("Draft generation", "Working locally", "Deterministic demo generator uses profile and freshest memory"),
        ("Edit and versioning", "Working", "Draft revisions increment the version and preserve status"),
        ("Approve and schedule", "Working locally", "Status changes persist and appear in Library/Calendar"),
        ("Profile settings", "Working", "Updates affect the next generated draft"),
        ("PostgreSQL schema", "Implemented", "SQLAlchemy models and Alembic initial migration"),
        ("Live LLM generation", "Simulated", "Anthropic/OpenAI providers remain placeholders"),
        ("LinkedIn OAuth/posting", "Simulated", "No real token exchange or publishing yet"),
        ("Authentication/payments", "Not integrated", "Existing API routes remain placeholders"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.05, 1.25, 3.2]
    headers = ["Capability", "Status", "Notes"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, PURPLE)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for capability, status, notes in rows:
        cells = table.add_row().cells
        values = [capability, status, notes]
        for index, value in enumerate(values):
            cells[index].width = Inches(widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
            cells[index].paragraphs[0].add_run(value)
        if status in ("Working", "Implemented"):
            set_cell_shading(cells[1], GREEN_LIGHT)
        elif status == "Working locally":
            set_cell_shading(cells[1], BLUE_LIGHT)
        else:
            set_cell_shading(cells[1], RED_LIGHT)
    set_table_borders(table)


def add_figure(doc, image_path, caption, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for style_name, size, before, after in (
    ("Heading 1", 17, 16, 8),
    ("Heading 2", 13.5, 12, 6),
    ("Heading 3", 11.5, 8, 4),
):
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(PURPLE if style_name != "Heading 3" else INK)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for style_name in ("List Bullet", "List Number"):
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.18

header = section.header.paragraphs[0]
header.text = "BLIDX WORKDESK  |  V1 IMPLEMENTATION REPORT"
header.style = styles["Normal"]
header.runs[0].font.size = Pt(8.5)
header.runs[0].font.bold = True
header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
footer = section.footer.paragraphs[0]
footer.add_run("Blidx - Internal MVP handoff  |  ")
footer.runs[0].font.size = Pt(8.5)
footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
add_page_number(footer)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(42)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(12)
run = kicker.add_run("PRODUCT + TECHNICAL WALKTHROUGH")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor.from_string(PURPLE)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
run = title.add_run("Blidx V1 Web App")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(32)
run.font.color.rgb = RGBColor.from_string(INK)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(22)
run = subtitle.add_run("Implementation Report and Test Guide")
run.font.size = Pt(17)
run.font.color.rgb = RGBColor.from_string(MUTED)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(26)
meta.add_run(f"Prepared by Jae  |  {date.today().strftime('%B %d, %Y')}  |  Local MVP").bold = True

add_callout(
    doc,
    "Purpose",
    "This document explains the first runnable Blidx web application built from the V1 System Design, Malia's integration guide, and the clickable HTML prototypes. It documents what can be tested now, what is simulated, and what remains for production.",
)
add_figure(doc, Path("/private/tmp/blidx-chat.png"), "Live local Blidx chat workspace with a personalized V2 draft.", 5.8)

add_page_break(doc)
add_heading(doc, "1. Executive Summary", 1)
doc.add_paragraph(
    "The current build converts the backend skeleton and UX prototypes into a working local web application. "
    "The goal of this milestone is not to claim a production-ready SaaS. It is to provide an honest, interactive "
    "product slice that demonstrates the core Blidx behavior: structured personal memory, personalized drafting, "
    "workflow ownership, approval, scheduling, and persistent content state."
)
add_callout(
    doc,
    "Core proof",
    "A founder can capture a real event in the Content Bank, ask Mira for a post, receive a draft that uses that event, revise it, approve it, and then see the result in the Library and Calendar.",
    GREEN_LIGHT,
    GREEN,
)
add_heading(doc, "What this milestone validates", 2)
for text in (
    "The product feels like a focused content workdesk rather than a general-purpose chatbot.",
    "Structured Content Bank entries can make drafts observably more personal.",
    "Draft state transitions can be represented clearly across Chat, Library, and Calendar.",
    "The frontend and backend can be tested as one integrated local application.",
    "The architecture can progress toward PostgreSQL, live LLM generation, and LinkedIn execution without changing the product concept.",
):
    add_bullet(doc, text)

add_heading(doc, "2. Implementation Status", 1)
add_status_table(doc)

add_page_break(doc)
add_heading(doc, "3. How to Run and Test", 1)
add_heading(doc, "Local startup", 2)
doc.add_paragraph("Open Terminal and run:")
code = doc.add_table(rows=1, cols=1)
code.alignment = WD_TABLE_ALIGNMENT.CENTER
code.autofit = False
code.columns[0].width = Inches(6.25)
cell = code.cell(0, 0)
set_cell_shading(cell, "F2F1F5")
set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
p = cell.paragraphs[0]
p.add_run(
    "cd /Users/kjysmu/Documents/Blidx\n"
    "source .venv/bin/activate\n"
    "uvicorn app.main:app --reload"
).font.name = "Courier New"
doc.add_paragraph("Open the following address in a browser:")
link = doc.add_paragraph()
r = link.add_run("http://127.0.0.1:8000/")
r.bold = True
r.font.color.rgb = RGBColor.from_string(PURPLE)

add_heading(doc, "Recommended demonstration sequence", 2)
steps = [
    "Open Content Bank and select a quick-capture category such as Met someone or Key insight.",
    "Enter a real founder moment and save it. Confirm that Blidx marks it Fresh.",
    "Return to Chat and request a draft about a relevant topic.",
    "Confirm that the new draft references the latest Content Bank event.",
    "Use Edit and request a bolder, shorter, or more personal version. Confirm that the version number increases.",
    "Approve the draft and choose Post now or Best time.",
    "Open Library to confirm the persisted status.",
    "Open Calendar to confirm the scheduled date is visible.",
    "Open Settings, change the tone or audience, and create another draft to observe personalization changes.",
]
for step in steps:
    add_number(doc, step)

add_heading(doc, "Automated verification", 2)
doc.add_paragraph("The backend and local product workflow are covered by automated tests:")
code2 = doc.add_table(rows=1, cols=1)
code2.alignment = WD_TABLE_ALIGNMENT.CENTER
code2.autofit = False
code2.columns[0].width = Inches(6.25)
cell = code2.cell(0, 0)
set_cell_shading(cell, "F2F1F5")
set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
cell.paragraphs[0].add_run("pytest").font.name = "Courier New"
doc.add_paragraph("Current result: 6 tests passing.")

add_page_break(doc)
add_heading(doc, "4. Product Walkthrough", 1)
add_heading(doc, "4.1 Chat and Draft Review", 2)
doc.add_paragraph(
    "Chat is the main operating surface. Mira displays weekly progress, Content Bank depth, and the current pending draft. "
    "The generated post uses profile data and the freshest Content Bank entry. Draft actions support approve, edit, save, and delete."
)
add_figure(doc, Path("/private/tmp/blidx-chat.png"), "Chat workspace: weekly goal, Content Bank depth, personalized draft, and review actions.")

add_heading(doc, "4.2 Content Bank", 2)
doc.add_paragraph(
    "The Content Bank captures the founder's real work and observations. Six quick categories mirror Malia's prototype: "
    "people, events, insights, milestones, reading, and solutions. Entries are stored as Fresh and surfaced in later drafts."
)
add_figure(doc, Path("/private/tmp/blidx-bank.png"), "Content Bank with quick-capture templates and persisted founder context.")

add_page_break(doc)
add_heading(doc, "4.3 Library", 2)
doc.add_paragraph(
    "The Library is the persistent content pipeline. It shows draft, scheduled, and published states, plus draft version and character count. "
    "Deleted drafts are soft-hidden from the user rather than erased from the demo store."
)
add_figure(doc, Path("/private/tmp/blidx-library.png"), "Library showing the revised post after scheduling.")

add_heading(doc, "4.4 Calendar", 2)
doc.add_paragraph(
    "The Calendar visualizes scheduled and published content. The local implementation uses the stored scheduled timestamp. "
    "Green denotes published content and purple denotes scheduled content."
)
add_figure(doc, Path("/private/tmp/blidx-calendar.png"), "Calendar view with the scheduled post visible on its target date.")

add_page_break(doc)
add_heading(doc, "4.5 Settings and Personalization", 2)
doc.add_paragraph(
    "Settings provides editable founder context: name, role, company, industry, description, audience, expertise, posting frequency, and tone. "
    "These values are read fresh when the next draft is created, matching the integration guide's expectation that profile changes affect Mira immediately."
)
add_figure(doc, Path("/private/tmp/blidx-settings.png"), "Settings panel used to control Mira's personalization context.")

add_heading(doc, "5. Current Architecture", 1)
doc.add_paragraph(
    "The current application remains a single FastAPI service. FastAPI serves both the web assets and the local MVP API. "
    "This minimizes setup friction while keeping a clean separation between interface state and backend behavior."
)
architecture = doc.add_table(rows=1, cols=3)
architecture.alignment = WD_TABLE_ALIGNMENT.CENTER
architecture.autofit = False
widths = [2.0, 2.2, 2.3]
for index, text in enumerate(("Layer", "Current implementation", "Production direction")):
    cell = architecture.rows[0].cells[index]
    cell.width = Inches(widths[index])
    set_cell_shading(cell, PURPLE)
    set_cell_margins(cell)
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
for values in (
    ("Frontend", "No-build mobile-first SPA served by FastAPI", "React/TypeScript PWA using the same API contracts"),
    ("API", "FastAPI `/api` demo endpoints", "Authenticated REST/SSE endpoints"),
    ("Persistence", "Local JSON demo store", "PostgreSQL via implemented SQLAlchemy/Alembic schema"),
    ("AI", "Deterministic personalized generator", "Claude primary, OpenAI failover, source-aware prompts"),
    ("Publishing", "Local status simulation", "LinkedIn OAuth and member posting API"),
):
    cells = architecture.add_row().cells
    for index, value in enumerate(values):
        cells[index].width = Inches(widths[index])
        cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cells[index])
        cells[index].paragraphs[0].add_run(value)
set_table_borders(architecture)

add_page_break(doc)
add_heading(doc, "6. Backend Behaviors Implemented", 1)
behaviors = [
    ("GET /api/state", "Returns the complete local user, profile, Content Bank, and posts state."),
    ("PUT /api/profile", "Updates founder profile and personalization fields."),
    ("POST /api/content-bank", "Stores a categorized entry with freshness and content-potential metadata."),
    ("POST /api/drafts", "Creates a personalized pending draft from profile and latest memory."),
    ("POST /api/drafts/{id}/edit", "Applies edit instructions, increments the version, and returns the draft to pending."),
    ("POST /api/drafts/{id}/approve", "Marks the post published locally or schedules it for the recommended time."),
    ("POST /api/drafts/{id}/save", "Moves the item into the Library as a saved draft."),
    ("POST /api/drafts/{id}/delete", "Soft-hides the item with deleted status."),
    ("POST /api/reset", "Resets local demo data for repeatable demonstrations."),
]
for endpoint, description in behaviors:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(endpoint)
    r.bold = True
    r.font.name = "Courier New"
    r.font.color.rgb = RGBColor.from_string(PURPLE)
    p.add_run(f" - {description}")

add_heading(doc, "State flow currently demonstrated", 2)
doc.add_paragraph("Pending draft -> Edited pending draft -> Scheduled or Published -> Visible in Library and Calendar.")
add_callout(
    doc,
    "Design note",
    "The integration guide specifies a richer draft state machine and permanent message history. The current slice implements only the states needed to validate the first end-to-end product journey.",
)

add_heading(doc, "7. Important Limitations", 1)
limitations = [
    "The local draft generator is deterministic and does not call Claude or OpenAI.",
    "There is no SSE activity stream yet; draft generation returns immediately.",
    "The local demo store is JSON, not the production PostgreSQL repository layer.",
    "Authentication, Google OAuth, password reset, Stripe checkout, and payment webhooks are not wired.",
    "LinkedIn OAuth and automatic publishing are not active. Post now marks the item published only inside the local app.",
    "Analytics, push notifications, email nudges, file uploads, URL extraction, and proactive cron jobs are not implemented.",
    "The web interface is a no-build SPA because the local machine currently has Node but no npm package manager. It should move to React/TypeScript before production UI work accelerates.",
]
for limitation in limitations:
    add_bullet(doc, limitation)

add_callout(
    doc,
    "Security action required",
    "The LinkedIn client secret appeared in the WhatsApp export. It must be rotated in the LinkedIn Developer Portal before any live integration. The exposed value was not copied into the project.",
    RED_LIGHT,
    "9B1C1C",
)

add_page_break(doc)
add_heading(doc, "8. Recommended Next Milestones", 1)
milestones = [
    ("Milestone 1 - Production persistence", "Implement repositories and switch profile, Content Bank, drafts, messages, and posts to PostgreSQL. Apply authentication ownership to every query."),
    ("Milestone 2 - Real authentication and onboarding", "Implement registration, login, JWT dependency, onboarding payload compatibility, and protected user-specific state."),
    ("Milestone 3 - Live Mira generation", "Add the model abstraction, Claude integration, system prompt, profile/memory context assembly, anti-hallucination rules, and structured draft output."),
    ("Milestone 4 - Chat history and SSE", "Persist messages and stream activity, message, and draft events according to Malia's integration guide."),
    ("Milestone 5 - LinkedIn connection and fallback", "Rotate credentials, implement OAuth, encrypt token storage, test member posting, and preserve the manual fallback experience."),
    ("Milestone 6 - React PWA conversion", "Move the validated UI into React/TypeScript components, add routing, PWA manifest/service worker, and stronger accessibility."),
    ("Milestone 7 - Proactive operation", "Add scheduling worker, content-gap detection, check-in prompts, notification records, and the one-action-per-user trigger lock."),
]
for title_text, detail in milestones:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title_text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(PURPLE)
    p.add_run(f"\n{detail}")

add_heading(doc, "9. Review Questions for Malia", 1)
for question in (
    "Does this first integrated product slice match the intended feeling of Mira owning the workflow?",
    "Is the Content Bank capture experience sufficiently clear for early founder tests?",
    "Should the next engineering priority be real AI generation or production authentication/onboarding?",
    "For early testers, is simulated LinkedIn publishing acceptable while OAuth approval and integration are completed?",
    "Which parts of the HTML prototypes are essential for the first founder test, and which can remain V1.1?",
):
    add_bullet(doc, question)

add_callout(
    doc,
    "Bottom line",
    "The current build is a functional local MVP demonstration, not a production launch. It is now suitable for internal review, workflow validation, and prioritizing the next engineering milestone with concrete product behavior in front of the team.",
    GREEN_LIGHT,
    GREEN,
)

doc.core_properties.title = "Blidx V1 Web App Implementation Report"
doc.core_properties.subject = "Product and technical walkthrough for the Blidx local MVP"
doc.core_properties.author = "Jae"
doc.core_properties.keywords = "Blidx, Mira, MVP, FastAPI, Content Bank, LinkedIn"
doc.save(OUT)
print(OUT)
